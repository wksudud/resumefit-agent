from io import BytesIO

from docx import Document

from src.resume_io import (
    build_rewritten_resume_doc,
    build_rewritten_resume_docx,
    build_rewritten_resume_markdown,
    read_uploaded_resume,
)
from src.schemas import RewriteSuggestion


class NamedBytesIO(BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


def test_reads_markdown_upload():
    uploaded = NamedBytesIO("# Resume\n\nPython project".encode("utf-8"), "resume.md")

    text, warning = read_uploaded_resume(uploaded)

    assert warning is None
    assert "Python project" in text


def test_reads_docx_upload():
    document = Document()
    document.add_paragraph("Resume Title")
    document.add_paragraph("Built an agent workflow.")
    buffer = BytesIO()
    document.save(buffer)
    uploaded = NamedBytesIO(buffer.getvalue(), "resume.docx")

    text, warning = read_uploaded_resume(uploaded)

    assert warning is None
    assert "Resume Title" in text
    assert "Built an agent workflow." in text


def test_legacy_doc_upload_extracts_best_effort_text():
    uploaded = NamedBytesIO(b"Resume Title\nBuilt an agent workflow.", "resume.doc")

    text, warning = read_uploaded_resume(uploaded)

    assert "Resume Title" in text
    assert warning == "legacy_doc_best_effort"


def test_legacy_doc_upload_warns_when_empty():
    uploaded = NamedBytesIO(b"\x00\x01\x02", "resume.doc")

    text, warning = read_uploaded_resume(uploaded)

    assert text == ""
    assert warning == "legacy_doc_empty"


def test_builds_rewritten_resume_markdown():
    suggestion = RewriteSuggestion(
        source_project="ResumeFit Agent",
        target_jd_requirement="Agent workflow",
        before_text="Built a tool.",
        after_text="Built a multi-agent resume matching workflow.",
        evidence="Project README",
    )

    markdown = build_rewritten_resume_markdown("Original resume", [suggestion])

    assert "Rewrite Suggestions" in markdown
    assert "Built a multi-agent resume matching workflow." in markdown


def test_builds_rewritten_resume_docx_bytes():
    docx_bytes = build_rewritten_resume_docx("# Draft\n\nResume body")

    assert docx_bytes.startswith(b"PK")


def test_builds_rewritten_resume_doc_bytes():
    doc_bytes = build_rewritten_resume_doc("# Draft\n\nResume body")

    assert doc_bytes.startswith(b"<html>")
    assert b"Resume body" in doc_bytes
