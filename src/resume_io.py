"""Resume upload parsing and rewritten resume export helpers."""

from __future__ import annotations

from html import escape
from io import BytesIO
from pathlib import Path
import re
from typing import BinaryIO

from docx import Document

from src.schemas import RewriteSuggestion


SUPPORTED_RESUME_TYPES = ("md", "markdown", "docx", "doc")


def read_uploaded_resume(uploaded_file) -> tuple[str, str | None]:
    """Extract text from a Streamlit uploaded resume file.

    Legacy .doc files are parsed on a best-effort basis because the binary Word
    format is not as reliable as Markdown or .docx in a cloud-only runtime.
    """
    if uploaded_file is None:
        return "", None

    file_name = getattr(uploaded_file, "name", "")
    suffix = Path(file_name).suffix.lower().lstrip(".")
    raw_bytes = _read_all(uploaded_file)

    if suffix in {"md", "markdown"}:
        return _decode_text(raw_bytes), None

    if suffix == "docx":
        document = Document(BytesIO(raw_bytes))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs), None

    if suffix == "doc":
        text = _extract_legacy_doc_text(raw_bytes)
        if text:
            return text, "legacy_doc_best_effort"
        return "", "legacy_doc_empty"

    return "", "unsupported"


def build_rewritten_resume_markdown(
    original_text: str,
    suggestions: list[RewriteSuggestion],
    language: str = "en",
) -> str:
    """Build a conservative rewritten-resume artifact from original text and suggestions."""
    labels = {
        "en": {
            "title": "Rewritten Resume Draft",
            "source": "Original Resume",
            "suggestions": "Rewrite Suggestions",
            "note": "Review each suggestion before using it in an application.",
            "before": "Before",
            "after": "After",
            "evidence": "Evidence",
        },
        "zh": {
            "title": "改写后简历草稿",
            "source": "原始简历",
            "suggestions": "改写建议",
            "note": "投递前请逐条核对事实与表述。",
            "before": "改写前",
            "after": "改写后",
            "evidence": "证据",
        },
    }.get(language, {})

    parts = [
        f"# {labels.get('title', 'Rewritten Resume Draft')}",
        f"> {labels.get('note', 'Review each suggestion before using it in an application.')}",
        "",
        f"## {labels.get('source', 'Original Resume')}",
        original_text.strip() or "_No resume text captured._",
        "",
        f"## {labels.get('suggestions', 'Rewrite Suggestions')}",
    ]

    if not suggestions:
        parts.append("_No rewrite suggestions generated._")
    else:
        for index, suggestion in enumerate(suggestions, 1):
            parts.extend([
                "",
                f"### {index}. {suggestion.source_project}",
                f"**{labels.get('before', 'Before')}:** {suggestion.before_text}",
                "",
                f"**{labels.get('after', 'After')}:** {suggestion.after_text}",
                "",
                f"**{labels.get('evidence', 'Evidence')}:** {suggestion.evidence}",
            ])

    return "\n".join(parts).strip() + "\n"


def build_rewritten_resume_docx(markdown_text: str) -> bytes:
    """Convert the generated Markdown draft into a simple Word document."""
    document = Document()

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("> "):
            document.add_paragraph(stripped[2:], style="Intense Quote")
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(_strip_basic_markdown(stripped))

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_rewritten_resume_doc(markdown_text: str) -> bytes:
    """Build a legacy .doc-compatible HTML document.

    Word can open HTML files saved with a .doc extension. This is more reliable
    in Streamlit Cloud than trying to write the old binary Word format directly.
    """
    body_lines: list[str] = []
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            body_lines.append(f"<h1>{escape(stripped[2:])}</h1>")
        elif stripped.startswith("## "):
            body_lines.append(f"<h2>{escape(stripped[3:])}</h2>")
        elif stripped.startswith("### "):
            body_lines.append(f"<h3>{escape(stripped[4:])}</h3>")
        elif stripped.startswith("> "):
            body_lines.append(f"<blockquote>{escape(stripped[2:])}</blockquote>")
        elif stripped.startswith("- "):
            body_lines.append(f"<p>&bull; {escape(stripped[2:])}</p>")
        else:
            body_lines.append(f"<p>{escape(_strip_basic_markdown(stripped))}</p>")

    html = "\n".join([
        "<html>",
        "<head>",
        '<meta charset="utf-8">',
        "<style>",
        "body { font-family: Calibri, Arial, sans-serif; line-height: 1.45; }",
        "h1, h2, h3 { color: #1f2937; }",
        "blockquote { color: #4b5563; border-left: 3px solid #9ca3af; padding-left: 12px; }",
        "</style>",
        "</head>",
        "<body>",
        *body_lines,
        "</body>",
        "</html>",
    ])
    return html.encode("utf-8")


def _read_all(uploaded_file) -> bytes:
    if hasattr(uploaded_file, "getvalue"):
        return uploaded_file.getvalue()
    if isinstance(uploaded_file, (bytes, bytearray)):
        return bytes(uploaded_file)
    if isinstance(uploaded_file, BinaryIO):
        return uploaded_file.read()
    return uploaded_file.read()


def _decode_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _extract_legacy_doc_text(raw_bytes: bytes) -> str:
    ole_text = _extract_doc_ole_text(raw_bytes)
    if ole_text:
        return ole_text
    return _extract_printable_text(raw_bytes)


def _extract_doc_ole_text(raw_bytes: bytes) -> str:
    try:
        import olefile
    except ImportError:
        return ""

    stream = BytesIO(raw_bytes)
    if not olefile.isOleFile(stream):
        return ""

    stream.seek(0)
    try:
        with olefile.OleFileIO(stream) as ole:
            parts: list[bytes] = []
            for stream_name in ("WordDocument", "1Table", "0Table"):
                if ole.exists(stream_name):
                    parts.append(ole.openstream(stream_name).read())
    except OSError:
        return ""

    return _extract_printable_text(b"\n".join(parts))


def _extract_printable_text(raw_bytes: bytes) -> str:
    candidates = [
        _decode_text(raw_bytes),
        raw_bytes.decode("utf-16-le", errors="ignore"),
    ]
    cleaned_candidates = [_clean_printable_chunks(candidate) for candidate in candidates]
    deduped = max(cleaned_candidates, key=_text_signal_score, default=[])
    return "\n\n".join(deduped[:120])


def _clean_printable_chunks(text: str) -> list[str]:
    chunks = re.findall(r"[\w\s\u4e00-\u9fff,.;:!?()/%+\-#@]{4,}", text)
    cleaned = []
    for chunk in chunks:
        line = re.sub(r"\s+", " ", chunk).strip()
        if line and sum(ch.isalnum() or "\u4e00" <= ch <= "\u9fff" for ch in line) >= 4:
            cleaned.append(line)
    return list(dict.fromkeys(cleaned))


def _text_signal_score(lines: list[str]) -> int:
    text = "\n".join(lines)
    return sum(ch.isalnum() or "\u4e00" <= ch <= "\u9fff" for ch in text)


def _strip_basic_markdown(text: str) -> str:
    return (
        text.replace("**", "")
        .replace("__", "")
        .replace("`", "")
    )
